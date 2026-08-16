from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "《人格牌》-当前版本完整玩法规则规格书_v2.1.md"
OUT_PATH = ROOT / "《人格牌》-当前版本完整玩法规则规格书_v2.1.docx"

NAVY = "16202A"
INK = "272A2E"
MUTED = "646A70"
GOLD = "A68045"
PALE_GOLD = "F3EBDD"
PALE_BLUE = "EEF2F5"
WHITE = "FFFFFF"
LINE = "D7D2C9"
RED = "873D3A"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, bottom=None, left=None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    if bottom:
        node = OxmlElement("w:bottom")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(bottom.get("size", 8)))
        node.set(qn("w:space"), str(bottom.get("space", 5)))
        node.set(qn("w:color"), bottom.get("color", GOLD))
        p_bdr.append(node)
    if left:
        node = OxmlElement("w:left")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(left.get("size", 18)))
        node.set(qn("w:space"), str(left.get("space", 10)))
        node.set(qn("w:color"), left.get("color", GOLD))
        p_bdr.append(node)


def set_run_font(run, east_asia="Microsoft YaHei", latin="Aptos") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, size: float, bold=False, color=INK, east_asia="Microsoft YaHei") -> None:
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_inline(paragraph, text: str, *, default_bold=False, default_color=INK) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = default_bold
        if part.startswith("**") and part.endswith("**"):
            part = part[2:-2]
            bold = True
        run = paragraph.add_run(part)
        set_run_font(run)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(default_color)


def add_num_pr(paragraph, num_id: int, level: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])


def create_abstract_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(4):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else f"%{level + 1}.")
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        p_pr.extend([tabs, ind])
        lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
        abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def new_num_id(doc: Document, abstract_id: int, start_at: int = 1) -> int:
    numbering = doc.part.numbering_part.element
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    for level in range(4):
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start_at if level == 0 else 1))
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.05)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, 10.2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.32
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    for style_name in ["List Number", "List Bullet"]:
        style = doc.styles[style_name]
        set_style_font(style, 10.0)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.widow_control = True

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, 18, bold=True, color=NAVY)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, 13.2, bold=True, color=NAVY)
    h2.paragraph_format.space_before = Pt(13)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, 11.0, bold=True, color=GOLD)
    h3.paragraph_format.space_before = Pt(9)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    if "Scope Note" not in doc.styles:
        style = doc.styles.add_style("Scope Note", WD_STYLE_TYPE.PARAGRAPH)
    scope = doc.styles["Scope Note"]
    set_style_font(scope, 9.6, color=MUTED)
    scope.paragraph_format.left_indent = Cm(0.55)
    scope.paragraph_format.right_indent = Cm(0.35)
    scope.paragraph_format.space_before = Pt(7)
    scope.paragraph_format.space_after = Pt(10)
    scope.paragraph_format.line_spacing = 1.3


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run("PERSONA CARD · RULES SPECIFICATION")
    set_run_font(r)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("人格牌")
    set_run_font(r)
    r.font.size = Pt(38)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("当前版本完整玩法规则规格书")
    set_run_font(r)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)
    set_paragraph_border(p, bottom={"color": GOLD, "size": 14, "space": 10})

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.45
    add_inline(p, "标准扑克牌组牌计分 × 人格牌构筑 × Boss 行为观察\n三场战斗、两次幕间选择、一次局终人格铸造", default_bold=True, default_color=INK)

    p = doc.add_paragraph()
    set_paragraph_shading(p, PALE_GOLD)
    set_paragraph_border(p, left={"color": GOLD, "size": 18, "space": 10})
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.right_indent = Cm(0.35)
    add_inline(p, "规则基线：当前网页 Demo\n文档版本：V2.1\n校准日期：2026-08-14\n用途：策划、程序、UI、QA 共同执行", default_color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("RULES AUTHORITY")
    set_run_font(r)
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "本文以当前可运行版本和自动测试结果为依据。旧版设想只有在本文明确保留时才继续有效。", default_color=MUTED)


def add_body_section(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.75)
    sec.bottom_margin = Cm(1.65)
    sec.left_margin = Cm(2.15)
    sec.right_margin = Cm(2.05)
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(4)
    r = hp.add_run("《人格牌》 · 当前版本完整玩法规则规格书 · V2.1")
    set_run_font(r)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    set_paragraph_border(hp, bottom={"color": LINE, "size": 5, "space": 4})

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("—  ")
    set_run_font(r)
    r.font.color.rgb = RGBColor.from_string(GOLD)
    add_field(fp, "PAGE", "1")
    r = fp.add_run("  —")
    set_run_font(r)
    r.font.color.rgb = RGBColor.from_string(GOLD)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.page_break_before = False
    add_inline(p, "目录", default_bold=True, default_color=NAVY)
    set_paragraph_border(p, bottom={"color": GOLD, "size": 10, "space": 7})

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    add_inline(intro, "目录将在 Word 中自动更新；标题可用于文档导航窗格。", default_color=MUTED)

    toc = doc.add_paragraph()
    toc.paragraph_format.line_spacing = 1.35
    add_field(toc, 'TOC \\o "1-1" \\h \\z \\u', "右键选择“更新域”以刷新目录")
    doc.add_page_break()


def style_heading(paragraph, level: int) -> None:
    if level == 1:
        set_paragraph_border(paragraph, bottom={"color": GOLD, "size": 8, "space": 6})
        # Small chapter marker in the right tab area is provided by the numbering in text.
    elif level == 2:
        set_paragraph_border(paragraph, left={"color": GOLD, "size": 12, "space": 7})


def parse_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    number_abs = create_abstract_numbering(doc, "number")
    bullet_abs = create_abstract_numbering(doc, "bullet")
    current_num = None
    current_bullet = None
    prior_list_kind = None
    in_frontmatter = True

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if in_frontmatter:
            if stripped.startswith("> "):
                in_frontmatter = False
            elif stripped.startswith("## "):
                in_frontmatter = False
            else:
                continue

        if not stripped:
            prior_list_kind = None
            current_num = None
            current_bullet = None
            continue

        if stripped.startswith("# "):
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:], default_bold=True, default_color=NAVY)
            style_heading(p, 1)
            prior_list_kind = None
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:], default_bold=True, default_color=NAVY)
            style_heading(p, 2)
            prior_list_kind = None
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:], default_bold=True, default_color=GOLD)
            prior_list_kind = None
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Scope Note")
            set_paragraph_shading(p, PALE_GOLD)
            set_paragraph_border(p, left={"color": GOLD, "size": 18, "space": 10})
            add_inline(p, stripped[2:], default_color=MUTED)
            prior_list_kind = None
            continue

        numbered = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        bullet = re.match(r"^(\s*)-\s+(.*)$", line)

        if numbered:
            indent, source_number, content = numbered.groups()
            level = min(3, len(indent) // 3)
            if prior_list_kind != "number":
                current_num = new_num_id(doc, number_abs, int(source_number))
            p = doc.add_paragraph(style="List Number")
            add_num_pr(p, current_num, level)
            add_inline(p, content)
            prior_list_kind = "number"
            continue

        if bullet:
            indent, content = bullet.groups()
            level = min(3, len(indent) // 2)
            if prior_list_kind != "bullet":
                current_bullet = new_num_id(doc, bullet_abs)
            p = doc.add_paragraph(style="List Bullet")
            add_num_pr(p, current_bullet, level)
            add_inline(p, content)
            prior_list_kind = "bullet"
            continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.72)
        if stripped.startswith("**") and stripped.endswith("**"):
            set_paragraph_shading(p, PALE_BLUE)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.right_indent = Cm(0.35)
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(7)
        add_inline(p, stripped)
        prior_list_kind = None


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    markdown = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "《人格牌》当前版本完整玩法规则规格书"
    doc.core_properties.subject = "当前网页 Demo 玩法规则权威稿"
    doc.core_properties.author = "项目策划组"
    doc.core_properties.keywords = "人格牌, 扑克牌, Roguelike, 规则规格书"
    doc.core_properties.comments = "V2.1 · 以 2026-08-14 当前 Demo 为规则基线"

    add_cover(doc)
    add_body_section(doc)
    add_toc(doc)
    parse_markdown(doc, markdown)
    set_update_fields_on_open(doc)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
